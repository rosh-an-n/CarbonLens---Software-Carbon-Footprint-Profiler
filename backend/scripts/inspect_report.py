import sys
from docx import Document

def extract_docx(file_path):
    try:
        doc = Document(file_path)
        print(f"--- Document: {file_path} ---")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[{i}] {para.text}")
        
        print("\n--- Tables ---")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}:")
            for row in table.rows:
                print(" | ".join(cell.text.strip() for cell in row.cells))
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    extract_docx("/Users/roshan/Projects/Software Carbon Footprint Profiler/CarbonLens_PBL_Report.docx")
