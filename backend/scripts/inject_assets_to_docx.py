import os
from docx import Document
from docx.shared import Inches

def update_report(docx_path, output_path, assets):
    doc = Document(docx_path)
    
    # helper to replace text in cell and add image
    def replace_with_image(table_index, placeholder_text, image_path, width=Inches(5.5)):
        if not os.path.exists(image_path):
            print(f"Warning: Image not found at {image_path}")
            return
            
        found = False
        table = doc.tables[table_index]
        for row in table.rows:
            for cell in row.cells:
                if placeholder_text in cell.text:
                    # Clear cell text
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=width)
                    found = True
                    break
            if found: break
        
        if not found:
            print(f"Placeholder '{placeholder_text}' not found in Table {table_index}")

    # Map our assets
    # Table indices were found during inspection
    # Table 2: Architecture
    replace_with_image(2, "Paste screenshot here: System Architecture", assets['architecture'], width=Inches(5.0))
    
    # Table 6: Dashboard
    replace_with_image(6, "Paste screenshot here: CarbonLens Dashboard", assets['dashboard'])
    
    # Table 7: Run Experiment
    replace_with_image(7, "Paste screenshot here: Run Experiment Page", assets['run_exp'])
    
    # Table 8: Results
    replace_with_image(8, "Paste screenshot here: Experiment Results", assets['results'])
    
    # Table 9: Compare
    replace_with_image(9, "Paste screenshot here: Compare Page", assets['compare'])
    
    # Table 10: History
    replace_with_image(10, "Paste screenshot here: History Page", assets['history'])
    
    # Table 11: CMT Submission
    replace_with_image(11, "Paste screenshot here: CMT Submission", assets['cmt'], width=Inches(5.0))
    
    # Save the updated document
    doc.save(output_path)
    print(f"Report updated successfully: {output_path}")

if __name__ == "__main__":
    brain_dir = "/Users/roshan/.gemini/antigravity/brain/2d08f93c-a6bc-4cb1-923d-e36eecc669c9"
    assets_map = {
        'architecture': f"{brain_dir}/carbonlens_architecture_1776483144863.png",
        'dashboard': f"{brain_dir}/dashboard_home_1776483632876.png",
        'run_exp': f"{brain_dir}/run_experiment_full_1776484141535.png",
        'results': f"{brain_dir}/run_experiment_full_1776484141535.png", # using same as it shows values
        'compare': f"{brain_dir}/compare_page_growth_curve_1776484278993.png",
        'history': f"{brain_dir}/history_page_attempt_1776484503606.png",
        'cmt': f"{brain_dir}/cmt_submission_screenshot_1776484769508.png",
    }
    
    update_report(
        "/Users/roshan/Projects/Software Carbon Footprint Profiler/CarbonLens_PBL_Report.docx",
        "/Users/roshan/Projects/Software Carbon Footprint Profiler/CarbonLens_PBL_Report_Updated.docx",
        assets_map
    )
